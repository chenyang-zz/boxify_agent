import logging
from io import BytesIO

import docx
import pymupdf
from bs4 import BeautifulSoup

log = logging.getLogger(__name__)


class DocumentParser:
    """Notebook文档解析领域服务，把不同文件格式统一转换为纯文本。"""

    SUPPORTED_EXTS = {".txt", ".md", ".markdown", ".html", ".htm", ".pdf", ".docx"}

    @classmethod
    def parse(cls, file_ext: str, content: bytes) -> str:
        """根据扩展名选择解析器，未知格式交由调用方处理为业务失败。"""
        ext = file_ext.lower()
        try:
            if ext in {".txt", ".md", ".markdown"}:
                return content.decode("utf-8", errors="ignore")
            if ext in {".html", ".htm"}:
                return BeautifulSoup(
                    content.decode("utf-8", errors="ignore"), "html.parser"
                ).get_text("\n")
            if ext == ".pdf":
                return cls._parse_pdf(content)
            if ext == ".docx":
                return cls._parse_docx(content)
        except UnicodeDecodeError as e:
            logging.error(f"文档解析失败: {str(e)}")
            raise

        raise ValueError(f"不支持的文件类型: {file_ext}")

    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        """从 PDF 页面中提取文本，保留页之间的换行边界。"""
        with pymupdf.Document(stream=content, filetype="pdf") as doc:
            return "\n".join(page.get_text() for page in doc)

    @staticmethod
    def _parse_docx(content: bytes) -> str:
        """从 DOCX 段落中提取文本，适配普通文档导入场景。"""
        document = docx.Document(BytesIO(content))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
